from pathlib import Path
import sys

from docx import Document


OLD_VALUE = "Pending owner approval"
NEW_VALUE = "1 September 2026"


def replace_in_paragraph(paragraph) -> int:
    replacements = 0
    for run in paragraph.runs:
        if OLD_VALUE in run.text:
            run.text = run.text.replace(OLD_VALUE, NEW_VALUE)
            replacements += 1
    return replacements


def main() -> int:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: correct_hr_policy_v1_effective_date.py INPUT.docx OUTPUT.docx")
    source = Path(sys.argv[1])
    destination = Path(sys.argv[2])
    document = Document(source)
    replacements = sum(replace_in_paragraph(paragraph) for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replacements += sum(replace_in_paragraph(paragraph) for paragraph in cell.paragraphs)
    if replacements != 1:
        raise RuntimeError(f"Expected exactly one effective-date replacement; found {replacements}.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    print(f"Corrected {replacements} metadata value in {destination}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
